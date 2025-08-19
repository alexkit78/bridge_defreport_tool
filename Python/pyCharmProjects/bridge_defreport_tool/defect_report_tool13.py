import sqlite3
import tkinter as tk
from tkinter import ttk, messagebox, filedialog
from docx import Document
from docx.shared import Pt
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
import shutil
import sys
import os
import ctypes
import uuid
from pathlib import Path


def resource_path(relative_path):
    if hasattr(sys, '_MEIPASS'):
        base_path = sys._MEIPASS
    else:
        base_path = os.path.abspath(".")
    return os.path.join(base_path, relative_path)


DB_PATH = resource_path("bridge_defects.db")
TEMPLATE_PATH = resource_path("report_template.docx")


class DefectApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Оценка состояния мостового сооружения")

        self.conn = sqlite3.connect(DB_PATH)
        self.cursor = self.conn.cursor()

        # report_data — список словарей, каждая запись имеет ключи:
        # placement, location, name, option, safety, durability, repairability, loadcap, action, uid
        self.report_data = []
        self.defect_options_by_numodm = {}
        self.defect_numodm_map = {}
        self.defect_categories_by_option = {}

        self.build_ui()

    def build_ui(self):
        frame = ttk.Frame(self.root, padding=10)
        frame.pack(fill="both", expand=True)

        frame.grid_columnconfigure(0, weight=1)
        frame.grid_columnconfigure(1, weight=1)

        style = ttk.Style()
        style.configure("Treeview", rowheight=26)

        ttk.Label(frame, text="Раздел:").grid(row=0, column=0, sticky="w", padx=6, pady=2)
        self.placement_cb = ttk.Combobox(frame, state="readonly")
        self.placement_cb.grid(row=1, column=0, columnspan=2, sticky="ew", padx=5, pady=2)
        self.placement_cb.bind("<<ComboboxSelected>>", self.load_defects)

        ttk.Label(frame, text="Местоположение дефекта:").grid(row=2, column=0, sticky="w", padx=6, pady=2)
        self.location_entry = ttk.Entry(frame)
        self.location_entry.grid(row=3, column=0, columnspan=2, sticky="ew", padx=5, pady=2)

        ttk.Label(frame, text="Поиск по типу дефекта:").grid(row=4, column=0, sticky="w", padx=6, pady=2)
        self.search_entry = ttk.Entry(frame)
        self.search_entry.grid(row=5, column=0, columnspan=2, sticky="ew", padx=5, pady=2)
        self.search_entry.bind("<KeyRelease>", self.filter_defect_names)

        ttk.Label(frame, text="Тип дефекта:").grid(row=6, column=0, sticky="w", padx=6, pady=2)
        self.defect_cb = ttk.Combobox(frame, state="readonly")
        self.defect_cb.grid(row=7, column=0, columnspan=2, sticky="ew", padx=5, pady=2)
        self.defect_cb.bind("<<ComboboxSelected>>", self.populate_defect_fields)

        ttk.Label(frame, text="Описание дефекта:").grid(row=8, column=0, sticky="w", padx=6, pady=2)
        self.option_cb = ttk.Combobox(frame, state="readonly")
        self.option_cb.grid(row=9, column=0, columnspan=2, sticky="ew", padx=5, pady=2)
        self.option_cb.bind("<<ComboboxSelected>>", self.populate_category_fields)

        ttk.Label(frame, text="Категории (Б, Д, Р, Г):").grid(row=10, column=0, sticky="w", padx=6, pady=2)
        cat_frame = ttk.Frame(frame)
        cat_frame.grid(row=11, column=0, columnspan=2, sticky="w")

        self.safety_entry = ttk.Entry(cat_frame, width=5)
        self.durability_entry = ttk.Entry(cat_frame, width=5)
        self.repair_entry = ttk.Entry(cat_frame, width=5)
        self.load_entry = ttk.Entry(cat_frame, width=5)

        self.safety_entry.pack(side="left", padx=5)
        self.durability_entry.pack(side="left", padx=5)
        self.repair_entry.pack(side="left", padx=5)
        self.load_entry.pack(side="left", padx=5)

        ttk.Label(frame, text="Мероприятия по устранению дефекта:").grid(row=12, column=0, sticky="w", padx=6, pady=2)
        self.action_entry = ttk.Entry(frame)
        self.action_entry.grid(row=13, column=0, columnspan=2, sticky="ew", padx=5, pady=2)

        ttk.Button(frame, text="Добавить в отчёт", command=self.add_entry).grid(row=14, column=0, pady=10)
        ttk.Button(frame, text="Сохранить в Word", command=self.export_docx).grid(row=14, column=1, pady=10)

        self.load_placements()

        # --- Таблица с прокруткой (Treeview + стандартные Scrollbar) ---
        table_container = ttk.Frame(self.root)
        table_container.pack(fill="both", expand=True)

        self.tree_scroll_y = ttk.Scrollbar(table_container, orient="vertical")
        self.tree_scroll_y.pack(side="right", fill="y")

        self.tree_scroll_x = ttk.Scrollbar(table_container, orient="horizontal")
        self.tree_scroll_x.pack(side="bottom", fill="x")

        self.table = ttk.Treeview(
            table_container,
            columns=("Раздел", "Местоположение", "Тип", "Описание", "Категории", "Мероприятия"),
            show="headings",
            yscrollcommand=self.tree_scroll_y.set,
            xscrollcommand=self.tree_scroll_x.set
        )

        # связать Scrollbar и Treeview
        self.tree_scroll_y.config(command=self.table.yview)
        self.tree_scroll_x.config(command=self.table.xview)

        for col in self.table["columns"]:
            self.table.heading(col, text=col)
            # разрешаем менять ширину столбца вручную
            self.table.column(col, minwidth=50, width=150, stretch=False)

        self.table.pack(fill="both", expand=True)

        # биндим двойной клик для редактирования и правый клик для контекстного меню
        self.table.bind("<Double-1>", self.start_cell_edit)
        self.table.bind("<Button-3>", self.show_context_menu)
        # self.table.bind("<Control-Button-1>", self.show_context_menu)
        self.table.bind("<Button-2>", self.show_context_menu) # пробуем
        # контекстное меню удаления строки по правому клику

        # Обработчики колесика мыши — на всякий случай обеспечим корректное скроллирование
        if sys.platform == 'darwin':
            # MacOS
            self.table.bind_all('<MouseWheel>', self._on_mousewheel_mac)
        else:
            # Windows / Linux
            self.table.bind_all('<MouseWheel>', self._on_mousewheel)
            self.table.bind_all('<Button-4>', lambda e: self.table.yview_scroll(-1, 'units'))
            self.table.bind_all('<Button-5>', lambda e: self.table.yview_scroll(1, 'units'))

        status_frame = ttk.Frame(self.root)
        status_frame.pack(side="bottom", fill="x", pady=3)
        self.status_label = ttk.Label(status_frame, text="",
                                      foreground="green", anchor="w")
        self.status_label.pack(side="left", padx=8)
        self.status_bar = ttk.Label(status_frame, text="", anchor="e")
        self.status_bar.pack(side="right", padx=8)
        self.update_status_bar()  # начальное обновление

        self.root.bind_class("Entry", "<Control-a>",
                             lambda e: e.widget.select_range(0, tk.END))
        self.root.bind_class("Entry", "<Command-a>",
                             lambda e: e.widget.select_range(0, tk.END))
        self.root.bind_class("Entry", "<Control-A>",
                             lambda e: e.widget.select_range(0, tk.END))

        # Вызов стандартного меню для Entry / Text / Combobox
        self.root.bind_class("Entry",
                             "<Button-3>" if sys.platform != 'darwin' else "<Button-2>",
                             self._show_entry_menu)
        self.root.bind_class("TEntry",
                             "<Button-3>" if sys.platform != 'darwin' else "<Button-2>",
                             self._show_entry_menu)
        self.root.bind_class("Text",
                             "<Button-3>" if sys.platform != 'darwin' else "<Button-2>",
                             self._show_entry_menu)
        self.root.bind_class("TCombobox",
                             "<Button-3>" if sys.platform != 'darwin' else "<Button-2>",
                             self._show_entry_menu)

        # Горячие клавиши в любой раскладке
        self.root.bind_all("<Control-KeyPress>", self._global_shortcuts)

    def _global_shortcuts(self, event):
        key = event.keysym.lower()
        widget = self.root.focus_get()
        if key == 'c':
            try:
                widget.event_generate("<<Copy>>")
            except:
                pass
        elif key == 'v':
            try:
                widget.event_generate("<<Paste>>")
            except:
                pass
        elif key == 'x':
            try:
                widget.event_generate("<<Cut>>")
            except:
                pass

    def _show_entry_menu(self, event):
        widget = event.widget
        menu = tk.Menu(self.root, tearoff=0)
        menu.add_command(label="Копировать",
                         command=lambda: widget.event_generate("<<Copy>>"))
        menu.add_command(label="Вставить",
                         command=lambda: widget.event_generate("<<Paste>>"))
        menu.add_command(label="Вырезать",
                         command=lambda: widget.event_generate("<<Cut>>"))
        menu.tk_popup(event.x_root, event.y_root)

    # --- скролл-обработчики ---
    def _on_mousewheel(self, event):
        # Windows: event.delta кратно 120
        self.table.yview_scroll(int(-1 * (event.delta / 120)), 'units')

    def _on_mousewheel_mac(self, event):
        # MacOS: event.delta небольшое значение
        self.table.yview_scroll(int(-1 * event.delta), 'units')

    # --- загрузка списков ---
    def load_placements(self):
        self.cursor.execute("SELECT DISTINCT name FROM placements")
        placements = [row[0] for row in self.cursor.fetchall()]
        placements_sorted = sorted(placements,
                                   key=lambda x: int(x.split('.', 1)[0]))
        self.placement_cb['values'] = placements_sorted

    def load_defects(self, event=None):
        placement = self.placement_cb.get()
        self.cursor.execute(
            "SELECT num_ODM, name, option, safetyClass, durabilityClass, repairabilityClass, loadCapacity FROM defect_types WHERE placement = ?",
            (placement,))
        rows = self.cursor.fetchall()

        self.defect_options_by_numodm.clear()
        self.defect_numodm_map.clear()
        self.defect_categories_by_option.clear()

        for num_odm, name, option, s, d, r, l in rows:
            self.defect_numodm_map[name] = (num_odm, placement)
            if num_odm not in self.defect_options_by_numodm:
                self.defect_options_by_numodm[num_odm] = []
            if option:
                if option not in self.defect_options_by_numodm[num_odm]:
                    self.defect_options_by_numodm[num_odm].append(option)
            if option:
                self.defect_categories_by_option[option] = (s, d, r, l)

        # сформировать значения combobox для типов
        self.defect_cb["values"] = sorted([name for name, (num, pl) in self.defect_numodm_map.items() if pl == placement])
        self.defect_cb.set("")
        self.option_cb.set("")
        self.safety_entry.delete(0, tk.END)
        self.durability_entry.delete(0, tk.END)
        self.repair_entry.delete(0, tk.END)
        self.load_entry.delete(0, tk.END)

    def filter_defect_names(self, event=None):
        text = self.search_entry.get().lower()
        placement = self.placement_cb.get()
        filtered = [name for name, (num, pl) in self.defect_numodm_map.items() if pl == placement and text in name.lower()]
        self.defect_cb["values"] = sorted(set(filtered))

    def populate_defect_fields(self, event=None):
        name = self.defect_cb.get()
        num_odm, _ = self.defect_numodm_map.get(name, (None, None))
        options = self.defect_options_by_numodm.get(num_odm, [])
        self.option_cb["values"] = options
        if options:
            self.option_cb.set(options[0])
            self.populate_category_fields()
        else:
            self.option_cb.set("")
            self.safety_entry.delete(0, tk.END)
            self.durability_entry.delete(0, tk.END)
            self.repair_entry.delete(0, tk.END)
            self.load_entry.delete(0, tk.END)

    def populate_category_fields(self, event=None):
        option = self.option_cb.get()
        s, d, r, l = self.defect_categories_by_option.get(option, ("", "", "", ""))
        self.safety_entry.delete(0, tk.END)
        if s is not None:
            self.safety_entry.insert(0, str(int(float(s))) if str(s).strip() != "" else "")
        self.durability_entry.delete(0, tk.END)
        if d is not None:
            self.durability_entry.insert(0, str(int(float(d))) if str(d).strip() != "" else "")
        self.repair_entry.delete(0, tk.END)
        if r is not None:
            self.repair_entry.insert(0, str(int(float(r))) if str(r).strip() != "" else "")
        self.load_entry.delete(0, tk.END)
        if l is not None:
            try:
                self.load_entry.insert(0, str(int(float(l))))
            except Exception:
                self.load_entry.insert(0, str(l))

    def update_status_bar(self):
        total = len(self.report_data)
        self.status_bar.config(text=f"Всего дефектов: {total}")
    def add_entry(self):
        placement = self.placement_cb.get()
        location = self.location_entry.get()
        name = self.defect_cb.get()
        option = self.option_cb.get()
        safety = self.safety_entry.get()
        durability = self.durability_entry.get()
        repairability = self.repair_entry.get()
        loadcap = self.load_entry.get()
        action_text = self.action_entry.get()

        if not placement or not name:
            messagebox.showerror("Ошибка", "Выберите раздел и тип дефекта")
            return

        uid = str(uuid.uuid4())
        rec = {
            'uid': uid,
            'placement': placement,
            'location': location,
            'name': name,
            'option': option,
            'safety': safety,
            'durability': durability,
            'repairability': repairability,
            'loadcap': loadcap,
            'action': action_text
        }
        self.report_data.append(rec)

        categories = []
        if safety:
            categories.append(f"Б{str(safety)}")
        if durability:
            categories.append(f"Д{str(durability)}")
        if repairability:
            categories.append(f"Р{str(repairability)}")
        try:
            if loadcap and int(loadcap) == 1:
                categories.append("Г")
        except ValueError:
            pass

        # вставляем строку с iid = uid, чтобы потом было удобно синхронизировать
        self.table.insert("", "end", iid=uid, values=(placement, location, name, option, ", ".join(categories), action_text))

        self.status_label.config(text="Строка добавлена в отчёт")
        self.root.after(2500, lambda: self.status_label.config(text=""))
        self.action_entry.delete(0, tk.END)

        self.update_status_bar()

    def start_cell_edit(self, event):
        # Редактируемые столбцы: 1=location, 2=name, 3=option, 5=action
        region = self.table.identify("region", event.x, event.y)
        if region != "cell":
            return
        row_id = self.table.identify_row(event.y)
        col = self.table.identify_column(event.x)
        if not row_id or not col:
            return
        col_index = int(col.replace('#', '')) - 1
        if col_index in (0, 4):
            return  # Раздел и Категории не редактируем

        bbox = self.table.bbox(row_id, col)
        if not bbox:
            return
        x, y, width, height = bbox

        value = self.table.item(row_id, 'values')[col_index]

        # tk.Entry вместо ttk.Entry
        entry = tk.Entry(self.table)
        entry.place(x=x, y=y, width=width, height=height)
        entry.insert(0, value)
        entry.focus()

        # Горячие клавиши
        entry.bind("<Control-a>", lambda e: entry.select_range(0, tk.END))
        entry.bind("<Control-c>", lambda e: (self.root.clipboard_clear(),
                                             self.root.clipboard_append(
                                                 entry.selection_get())))
        entry.bind("<Control-v>", lambda e: entry.insert(tk.INSERT,
                                                         self.root.clipboard_get()))
        entry.bind("<Control-x>", lambda e: (self.root.clipboard_clear(),
                                             self.root.clipboard_append(
                                                 entry.selection_get()),
                                             entry.delete("sel.first",
                                                          "sel.last")))

        # Контекстное меню
        menu = tk.Menu(self.root, tearoff=0)
        menu.add_command(label="Копировать",
                         command=lambda: (self.root.clipboard_clear(),
                                          self.root.clipboard_append(
                                              entry.selection_get())))
        menu.add_command(label="Вставить",
                         command=lambda: entry.insert(tk.INSERT,
                                                      self.root.clipboard_get()))
        menu.add_command(label="Вырезать",
                         command=lambda: (self.root.clipboard_clear(),
                                          self.root.clipboard_append(
                                              entry.selection_get()),
                                          entry.delete("sel.first",
                                                       "sel.last")))

        entry.bind("<Button-3>",
                   lambda e: menu.tk_popup(e.x_root, e.y_root))  # только ПКМ

        def save_edit(event=None):
            new_value = entry.get()
            values = list(self.table.item(row_id, 'values'))
            values[col_index] = new_value
            self.table.item(row_id, values=values)
            entry.destroy()

            # синхронизация report_data
            for rec in self.report_data:
                if rec['uid'] == row_id:
                    if col_index == 1:
                        rec['location'] = new_value
                    elif col_index == 2:
                        rec['name'] = new_value
                    elif col_index == 3:
                        rec['option'] = new_value
                    elif col_index == 5:
                        rec['action'] = new_value
                    break
            self.update_status_bar()

        def cancel_edit(event=None):
            entry.destroy()

        entry.bind("<Return>", save_edit)
        entry.bind("<FocusOut>", save_edit)
        entry.bind("<Escape>", cancel_edit)


    def show_context_menu(self, event):
        row_id = self.table.identify_row(event.y)
        if row_id:
            self.table.selection_set(row_id)
            menu = tk.Menu(self.root, tearoff=0)
            menu.add_command(label="Удалить", command=self.delete_selected_row)
            try:
                menu.tk_popup(event.x_root, event.y_root)
            finally:
                menu.grab_release()



    def delete_selected_row(self):
        sel = self.table.selection()
        if not sel:
            return
        for iid in sel:
            # удалить из report_data
            self.report_data = [r for r in self.report_data if r['uid'] != iid]
            self.table.delete(iid)
        self.update_status_bar()

    def export_docx(self):
        if not self.report_data:
            messagebox.showwarning("Нет данных",
                                   "Сначала добавьте данные в отчёт.")
            return

        file_path = filedialog.asksaveasfilename(defaultextension=".docx",
                                                 filetypes=[("Word document",
                                                             "*.docx")])
        if not file_path:
            return

        try:
            shutil.copyfile(TEMPLATE_PATH, file_path)
            doc = Document(file_path)
        except Exception:
            doc = Document()
            table = doc.add_table(rows=1, cols=6)
        else:
            table = None
            expected_headers = ["№", "Местоположение", "Тип", "Описание",
                                "Категории", "Мероприятия"]
            for t in doc.tables:
                first_row_texts = [c.text.strip().lower() for c in
                                   t.rows[0].cells]
                if all(any(h.lower() in ct for ct in first_row_texts) for h in
                       expected_headers):
                    table = t
                    break
            if table is None:
                table = doc.tables[0]

        # очищаем все строки после заголовка
        hdr_rows = 1
        try:
            while len(table.rows) > hdr_rows:
                table._tbl.remove(table.rows[hdr_rows]._tr)
        except Exception:
            pass

        # группируем дефекты по placement
        from collections import defaultdict
        grouped = defaultdict(list)
        for rec in self.report_data:
            grouped[rec['placement']].append(rec)

        # сортируем placement по числовому префиксу
        placements_sorted = sorted(grouped.keys(),
                                   key=lambda x: int(x.split('.', 1)[0]))
        counter = 1
        
        for placement in placements_sorted:
            entries = grouped[placement]

            # строка-заголовок раздела
            placement_row = table.add_row()
            first_cell = placement_row.cells[0]
            last_cell = placement_row.cells[-1]
            first_cell.merge(last_cell)

            clean_placement = placement.split('.', 1)[
                -1].strip()  # убираем номер
            para = first_cell.paragraphs[0]
            para.text = ""
            run = para.add_run(clean_placement)
            run.bold = True
            run.font.size = Pt(12)
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            try:
                first_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            except Exception:
                pass

            # строки с дефектами
            for rec in entries:
                row_cells = table.add_row().cells
                row_cells[0].text = str(counter)  # номер строки
                row_cells[1].text = rec.get('location', '')
                row_cells[2].text = rec.get('name', '')
                row_cells[3].text = rec.get('option', '')

                cats = []
                if rec.get('safety'):
                    cats.append(f"Б{rec.get('safety')}")
                if rec.get('durability'):
                    cats.append(f"Д{rec.get('durability')}")
                if rec.get('repairability'):
                    cats.append(f"Р{rec.get('repairability')}")
                try:
                    if rec.get('loadcap') and int(rec.get('loadcap')) == 1:
                        cats.append("Г")
                except Exception:
                    pass

                row_cells[4].text = ", ".join(cats)
                row_cells[5].text = rec.get('action', '')

                # центрируем все ячейки
                for cell in row_cells:
                    try:
                        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    except Exception:
                        pass
                    for p in cell.paragraphs:
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

                counter += 1

        doc.save(file_path)
        messagebox.showinfo("Готово", f"Файл сохранён: {file_path}")


if __name__ == "__main__":
    if sys.platform.startswith('win'):
        try:
            ctypes.windll.shcore.SetProcessDpiAwareness(1)  # Windows 8.1+
        except:
            ctypes.windll.user32.SetProcessDPIAware()  # Windows 7
    root = tk.Tk()
    app = DefectApp(root)
    root.mainloop()
