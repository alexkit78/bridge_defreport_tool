# export.py
import shutil
import re
import os
from collections import defaultdict

from docx import Document
from docx.document import Document as _Document
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.shared import Pt, Cm
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from copy import deepcopy
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from constants import TEMPLATE_PATH, REPORT_TEMPLATE_PATH
from dictionary import BRIDGE_KEYS

FLOAT_1 = {"hydro_B", "hydro_H", "pier_height"}
FLOAT_2 = {
    "hydro_V", "under_clearance", "length", "opening",
    "width_B", "width_G", "width_C1", "width_C2", "width_T1", "width_T2",
    "guardrails_height_bridge", "guardrails_height_approach",
    "railings_height",
    "approach_width1", "approach_width2",
    "mound_height1", "mound_height2", "span_width_B", "span_width_G", "span_width_C1", "span_width_C2", "span_width_T1", "span_width_T2",
    "deck_thickness", "pavement_thickness",
    "main_beam_h_mid", "main_beam_h_support", "pier_size_a", "pier_size_b", "piles_spacing",
    "pier_rigel_width", "pier_rigel_height", "pier_rigel_length", "pavement_extrathickness",
}


# ==================================================
# Word helpers: порядок как в документе
# ==================================================

def iter_block_items(parent):
    """
    Идём по документу по порядку: абзацы и таблицы,
    в том виде, как они реально идут в Word.
    """
    parent_elm = parent.element.body
    for child in parent_elm.iterchildren():
        if child.tag.endswith('}p'):
            yield Paragraph(child, parent)
        elif child.tag.endswith('}tbl'):
            yield Table(child, parent)



def find_paragraph_with_marker(doc: Document, marker: str) -> Paragraph | None:
    for p in doc.paragraphs:
        if marker in p.text:
            return p
    return None

def insert_element_after(reference, new_element):
    """
    reference может быть:
    - Paragraph (тогда есть ._element)
    - XML элемент CT_P/CT_Tbl (тогда addnext прямо на нём)
    """
    ref_elm = getattr(reference, "_element", reference)
    ref_elm.addnext(new_element)

def _is_paragraph_elm(elm) -> bool:
    return elm.tag == qn("w:p")

def _is_table_elm(elm) -> bool:
    return elm.tag == qn("w:tbl")

def find_body_element_index_by_text(doc: Document, marker: str) -> int | None:
    body = doc.element.body
    for i, el in enumerate(list(body)):
        if _is_paragraph_elm(el):
            p = Paragraph(el, doc)
            if marker in p.text:
                return i
    return None

def replace_placeholders_in_body_slice(doc: Document, body_elements: list, mapping: dict):
    """
    Заменяет плейсхолдеры ТОЛЬКО в указанных элементах body (абзацы/таблицы).
    """
    for el in body_elements:
        if _is_paragraph_elm(el):
            replace_in_paragraph(Paragraph(el, doc), mapping)
        elif _is_table_elm(el):
            replace_in_table(Table(el, doc), mapping)

def clone_block_between_markers(
    doc: Document,
    start_marker: str,
    end_marker: str,
    items: list,
    prefix: str,
    mapping_builder=None
):
    """
    Клонирует блок документа, который находится между start_marker и end_marker
    (start_marker и end_marker — это абзацы с этими текстами).
    
    На каждый item создаётся своя копия блока,
    и плейсхолдеры {{prefix.*}} заменяются только внутри этой копии.
    """
    if not items:
        # если данных нет — просто уберём маркер, чтобы не торчал
        remove_marker_everywhere(doc, start_marker)
        return

    body = doc.element.body
    body_list = list(body)

    start_i = find_body_element_index_by_text(doc, start_marker)
    if start_i is None:
        raise ValueError(f"Не найден маркер {start_marker}")

    end_i = find_body_element_index_by_text(doc, end_marker)
    if end_i is None:
        raise ValueError(f"Не найден маркер {end_marker}")

    if end_i <= start_i:
        raise ValueError(f"Маркер {end_marker} должен быть после {start_marker}")

    # Блок формы = элементы между маркерами:
    # включаем абзац с start_marker (чтобы потом его удалить)
    # но end_marker НЕ включаем
    block = body_list[start_i:end_i]

    # Сначала удалим исходный блок
    # (удалять надо из body, не из body_list)
    for el in block:
        body.remove(el)

    # Теперь вставим N копий на место удаления, в том же порядке
    insert_pos = start_i  # куда вставлять в body

    for item in items:
        cloned_block = [deepcopy(el) for el in block]

        # вставляем в body по порядку
        for el in cloned_block:
            body.insert(insert_pos, el)
            insert_pos += 1

                
        replace_placeholders_in_body_slice(doc, cloned_block, {start_marker: ""})

        if mapping_builder:
            mapping = mapping_builder(item)
        else:
            mapping = {f"{{{{{prefix}.{k}}}}}": v for k,v in item.items() if k!="uid"}
        replace_placeholders_in_body_slice(doc, cloned_block, mapping)

    # В итоге в документе не осталось исходного start_marker, потому что он был внутри block


def find_table_after_marker(doc: Document, marker: str) -> Table | None:
    """
    Находит таблицу, которая идёт СРАЗУ ПОСЛЕ абзаца с marker.
    """
    seen_marker = False
    for block in iter_block_items(doc):
        if isinstance(block, Paragraph):
            if marker in block.text:
                seen_marker = True
        elif isinstance(block, Table):
            if seen_marker:
                return block
    return None


def remove_marker_everywhere(doc: Document, marker: str):
    # основной текст
    for p in doc.paragraphs:
        for run in p.runs:
            if marker in run.text:
                run.text = run.text.replace(marker, "")

    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        if marker in run.text:
                            run.text = run.text.replace(marker, "")

    # колонтитулы всех типов
    for section in doc.sections:
        headers = [section.header, section.first_page_header, section.even_page_header]
        footers = [section.footer, section.first_page_footer, section.even_page_footer]

        for hdr in headers:
            for p in hdr.paragraphs:
                for run in p.runs:
                    if marker in run.text:
                        run.text = run.text.replace(marker, "")
            for t in hdr.tables:
                for row in t.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            for run in p.runs:
                                if marker in run.text:
                                    run.text = run.text.replace(marker, "")

        for ftr in footers:
            for p in ftr.paragraphs:
                for run in p.runs:
                    if marker in run.text:
                        run.text = run.text.replace(marker, "")
            for t in ftr.tables:
                for row in t.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            for run in p.runs:
                                if marker in run.text:
                                    run.text = run.text.replace(marker, "")

def _safe_join(folder: str, filename: str) -> str:
    if not folder or not filename:
        return ""
    return os.path.join(folder, filename)

def _insert_picture_at_marker(doc: Document, marker: str, image_path: str, w_cm: float, h_cm: float):
    """
    Находит абзац с marker, удаляет marker, и вставляет картинку в этот абзац.
    Если картинки нет — просто убирает marker (остаётся пустое место).
    """
    p = find_paragraph_with_marker(doc, marker)
    if p is None:
        return

    # убираем маркер из текста
    p.text = p.text.replace(marker, "").strip()

    if not image_path or not os.path.isfile(image_path):
        return  # оставляем пустое место

    run = p.add_run()
    inline = run.add_picture(image_path, width=Cm(w_cm), height=Cm(h_cm))
    _ensure_effect_extent(inline, border_width_pt=0.75)
    _add_picture_border_inline(inline, width_pt=0.75)

def _add_picture_border_inline(inline, width_pt=0.75, color_hex="000000"):
    """
    Добавляет чёрную рамку вокруг InlineShape (картинки).
    inline: объект, который возвращает run.add_picture() (InlineShape)
    """
    width_emu = int(width_pt * 12700)

    try:
        inline_elm = inline._inline  # wp:inline
    except Exception:
        return

    # ищем pic:pic без namespaces, чтобы не ловить ошибки
    pics = inline_elm.xpath(".//*[local-name()='pic']")
    if not pics:
        return
    pic = pics[0]

    # внутри pic ищем spPr
    spPr = None
    for child in pic.iterchildren():
        if child.tag.endswith("}spPr"):
            spPr = child
            break

    if spPr is None:
        spPr = OxmlElement("pic:spPr")
        pic.append(spPr)

    # удалить старые линии (если были)
    for old_ln in spPr.xpath("./*[local-name()='ln']"):
        spPr.remove(old_ln)

    ln = OxmlElement("a:ln")
    ln.set("w", str(width_emu))
    ln.set("algn", "in")

    solidFill = OxmlElement("a:solidFill")
    srgbClr = OxmlElement("a:srgbClr")
    srgbClr.set("val", color_hex)
    solidFill.append(srgbClr)
    ln.append(solidFill)

    prstDash = OxmlElement("a:prstDash")
    prstDash.set("val", "solid")
    ln.append(prstDash)

    ln.append(OxmlElement("a:round"))
    spPr.append(ln)

def _pt_to_emu(pt: float) -> int:
    # 1 pt = 12700 EMU
    return int(round(pt * 12700))

def _ensure_effect_extent(inline, border_width_pt: float = 0.75):
    """
    Добавляет/обновляет wp:effectExtent у wp:inline,
    чтобы Picture Border не клиппился (особенно сверху).
    """
    # ВАЖНО:
    # Word клиппит picture border без wp:effectExtent.
    # Без этого рамка может пропадать сверху (и в PDF тоже).
    # Не удалять _ensure_effect_extent!
    try:
        inline_elm = inline._inline  # wp:inline
    except Exception:
        return

    pad = _pt_to_emu(border_width_pt)

    # ищем effectExtent без namespaces (python-docx иногда ругается на namespaces)
    eff = inline_elm.xpath(".//*[local-name()='effectExtent']")
    if eff:
        eff = eff[0]
    else:
        eff = OxmlElement("wp:effectExtent")

        # вставить по порядку: после wp:extent
        extent = inline_elm.xpath("./*[local-name()='extent']")
        if extent:
            extent = extent[0]
            extent.addnext(eff)
        else:
            inline_elm.insert(0, eff)

    eff.set("l", str(pad))
    eff.set("t", str(pad))
    eff.set("r", str(pad))
    eff.set("b", str(pad))

def _fill_photos_gallery(doc: Document, marker: str, photos: list, folder: str, w_cm: float, h_cm: float):
    """
    Вставляет блок фотографий:
    - всё по центру
    - размер шрифта подписи 14 pt
    - подпись ПОД фото: "Фото N. Описание"
    - space after paragraph, чтобы между блоками был промежуток
    """
    p = find_paragraph_with_marker(doc, marker)
    if p is None:
        return

    # Если фото нет — просто убираем маркер
    if not photos:
        p.text = p.text.replace(marker, "")
        return

    # Удаляем маркер-текст
    p.text = p.text.replace(marker, "").strip()

    # Вставляем элементы ПОСЛЕ этого абзаца
    ref = p._element

    num = 1
    for rec in photos:
        filename = (rec or {}).get("filename", "") or ""
        caption = (rec or {}).get("caption", "") or ""
        img_path = _safe_join(folder, filename)

        # --- 1) абзац с картинкой (по центру) ---
        pic_p = p._parent.add_paragraph()
        pic_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf = pic_p.paragraph_format
        pf.line_spacing_rule = WD_LINE_SPACING.SINGLE 
        pf.space_after = Pt(4)

        if img_path and os.path.isfile(img_path):
            run = pic_p.add_run()
            inline = run.add_picture(img_path, width=Cm(w_cm), height=Cm(h_cm))
            _ensure_effect_extent(inline, border_width_pt=0.75)
            _add_picture_border_inline(inline, width_pt=0.75)

        insert_element_after(ref, pic_p._element)
        ref = pic_p._element

        # --- 2) подпись под фото (по центру, 14 pt) ---
        cap_p = p._parent.add_paragraph()
        cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf = cap_p.paragraph_format
        pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
        pf.space_after = Pt(8)

        r = cap_p.add_run(f"Фото {num}. {caption}".strip())
        r.font.size = Pt(14)

        insert_element_after(ref, cap_p._element)
        ref = cap_p._element

        num += 1


# ==================================================
# PLACEHOLDERS
# ==================================================

def _copy_run_format(src_run, dst_run):
    """
    Копирует основные свойства форматирования run.
    Этого обычно достаточно, чтобы сохранить bold и размер шрифта.
    """
    dst_run.bold = src_run.bold
    dst_run.italic = src_run.italic
    dst_run.underline = src_run.underline

    # font settings
    dst_run.font.name = src_run.font.name
    dst_run.font.size = src_run.font.size
    dst_run.font.all_caps = src_run.font.all_caps
    dst_run.font.small_caps = src_run.font.small_caps
    dst_run.font.strike = src_run.font.strike
    dst_run.font.double_strike = src_run.font.double_strike
    dst_run.font.subscript = src_run.font.subscript
    dst_run.font.superscript = src_run.font.superscript

    # цвет иногда важен (если задан)
    if src_run.font.color and src_run.font.color.rgb:
        dst_run.font.color.rgb = src_run.font.color.rgb


def replace_in_paragraph(paragraph: Paragraph, mapping: dict):
    if not paragraph.runs:
        return

    full_text = "".join(run.text for run in paragraph.runs)
    new_text = full_text

    has_any_key = any(k in new_text for k in mapping.keys())
    if not has_any_key:
        return

    for key, val in mapping.items():
        new_text = new_text.replace(key, "" if val is None else str(val))

    new_text = _format_units_for_docx(new_text)

    if new_text == full_text:
        return

    # --- ВАЖНО: если где-то в абзаце была жирность, сохраним её ---
    bold_flag = any((r.bold is True) for r in paragraph.runs)

    # выберем run-образец (для размера/шрифта и т.п.)
    sample_run = paragraph.runs[0]

    # очищаем все runs
    for run in paragraph.runs:
        run.text = ""

    # записываем результат
    if paragraph.runs:
        paragraph.runs[0].text = new_text
        _copy_run_format(sample_run, paragraph.runs[0])

        # ключевое: вернуть жирность, если она была в абзаце
        if bold_flag:
            paragraph.runs[0].bold = True
    else:
        r = paragraph.add_run(new_text)
        _copy_run_format(sample_run, r)
        if bold_flag:
            r.bold = True


def replace_in_table(table: Table, mapping: dict):
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                replace_in_paragraph(p, mapping)



def replace_placeholders_everywhere(doc: _Document, mapping: dict):
    """
    Заменяет плейсхолдеры:
    - в тексте
    - в таблицах
    - во ВСЕХ вариантах колонтитулов (обычный / первый лист / чётные)
    """
    # основной текст
    for p in doc.paragraphs:
        replace_in_paragraph(p, mapping)

    for t in doc.tables:
        replace_in_table(t, mapping)

    # колонтитулы всех типов
    for section in doc.sections:
        headers = [
            section.header,
            section.first_page_header,
            section.even_page_header,
        ]
        footers = [
            section.footer,
            section.first_page_footer,
            section.even_page_footer,
        ]

        for hdr in headers:
            for p in hdr.paragraphs:
                replace_in_paragraph(p, mapping)
            for t in hdr.tables:
                replace_in_table(t, mapping)

            replace_placeholders_in_element_xml(hdr._element, mapping)

        for ftr in footers:
            for p in ftr.paragraphs:
                replace_in_paragraph(p, mapping)
            for t in ftr.tables:
                replace_in_table(t, mapping)
            replace_placeholders_in_element_xml(ftr._element, mapping)

def replace_placeholders_in_element_xml(element, mapping: dict):
    # без namespace-префиксов, чтобы работало всегда
    for p in element.xpath(".//*[local-name()='p']"):
        ts = p.xpath(".//*[local-name()='t']")
        if not ts:
            continue

        full_text = "".join((t.text or "") for t in ts)
        if not any(k in full_text for k in mapping.keys()):
            continue

        new_text = full_text
        for key, val in mapping.items():
            new_text = new_text.replace(key, "" if val is None else str(val))

        if new_text == full_text:
            continue

        ts[0].text = new_text
        for t in ts[1:]:
            t.text = ""
# ==================================================
# TEXT NORMALIZATION
# ==================================================
def _normalize_dash(text: str) -> str:
    """
    1) если введено ровно "-" -> "—"
    2) если есть " - " -> " — " (для красивых названий)
    """
    if text is None:
        return ""
    s = str(text).strip()
    if s == "-":
        return "–"
    # заменяем только дефис-разделитель с пробелами
    return str(text).replace(" - ", " – ")

def _format_units_for_docx(text: str) -> str:
    if not text:
        return text
    return (
        text
        .replace("м2", "м²")
        .replace("м3", "м³")
    )

def _keep_highlight_if_empty(s: str) -> str:
    # NBSP чтобы подсветка маркером (заливка) визуально оставалась
    return s if (s is not None and str(s).strip() != "") else "\u00A0"

def _yes_no_to_10(value: str) -> str:
    v = (value or "").strip().lower()
    if v == "да":
        return "1"
    if v == "нет":
        return "0"
    return ""

def _flow_dir_to_sign(value: str) -> str:
    v = (value or "").strip().lower()
    if v == "слева направо":
        return "1"
    if v == "справа налево":
        return "-1"
    return ""

def _fmt_float(value: str, decimals: int) -> str:
    s = (value or "").strip().replace(",", ".")
    if s == "":
        return "\u00A0"
    try:
        x = float(s)
    except ValueError:
        return _keep_highlight_if_empty(value)
    return f"{x:.{decimals}f}".replace(".", ",")

def _calc_km_code(km_value: str) -> str:
    """
    Из '9+700'  -> '010'
    Из '23,245' -> '024'
    Из '325+000' -> '326'
    Из '1251+...' -> '1252' для кода сооружения
    """
    if not km_value:
        return ""

    s = str(km_value).strip()

    for sep in ("+", ","):
        if sep in s:
            left = s.split(sep, 1)[0].strip()
            try:
                km = int(left) + 1
            except ValueError:
                return ""

            return f"{km:03d}" if km < 1000 else str(km)

    return ""


def prepare_bridge_mapping(bridge: dict) -> dict:
    """
    Делает mapping для плейсхолдеров {{bridge.*}} с нужными преобразованиями.
    """
    b = bridge or {}

    mapping = {}
    for k in BRIDGE_KEYS:
        v = b.get(k, "")
        # базовая строка
        sv = "" if v is None else str(v)

        # дефис/тире (применяем для всех — безопасно)
        sv = _normalize_dash(sv)

        if k in FLOAT_1:
            sv = _fmt_float(sv, 1)
        elif k in FLOAT_2:
            sv = _fmt_float(sv, 2)

        # точечные конвертации
        if k in ("marking", "transition_slabs"):
            sv = _yes_no_to_10(sv)

        if k == "flow_direction":
            sv = _flow_dir_to_sign(sv)

        mapping[f"{{{{bridge.{k}}}}}"] = _keep_highlight_if_empty(sv)
        km_code = _calc_km_code(b.get("km", ""))
        mapping["{{bridge.km_code}}"] = _keep_highlight_if_empty(km_code)

    return mapping

def prepare_bridge_mapping_report(bridge: dict) -> dict:
    """
    Mapping для ТЕХНИЧЕСКОГО ОТЧЁТА.
    Отличие от паспорта:
    - flow_direction вставляется ТЕКСТОМ, а не числом
    """
    b = bridge or {}
    mapping = {}

    for k in BRIDGE_KEYS:
        v = b.get(k, "")
        sv = "" if v is None else str(v)

        # тире / дефисы
        sv = _normalize_dash(sv)

        # формат чисел
        if k in FLOAT_1:
            sv = _fmt_float(sv, 1)
        elif k in FLOAT_2:
            sv = _fmt_float(sv, 2)

        # да/нет → 1/0 (оставляем, это в отчёте тоже ок)
        if k in ("marking", "transition_slabs"):
            sv = _yes_no_to_10(sv)

        # !!! ВАЖНО !!!
        # flow_direction НЕ преобразуем в знак
        # оставляем текст: "слева направо", "справа налево"

        mapping[f"{{{{bridge.{k}}}}}"] = _keep_highlight_if_empty(sv)

    # km_code тоже нужен
    km_code = _calc_km_code(b.get("km", ""))
    mapping["{{bridge.km_code}}"] = _keep_highlight_if_empty(km_code)

    return mapping

def prepare_span_mapping(span: dict) -> dict:
    s = span or {}
    mapping = {}

    for k, v in s.items():
        if k == "uid":
            continue
        sv = "" if v is None else str(v)
        sv = _normalize_dash(sv)


        if k in FLOAT_1:
            sv = _fmt_float(sv, 1)
        elif k in FLOAT_2:
            sv = _fmt_float(sv, 2)

        mapping[f"{{{{span.{k}}}}}"] = _keep_highlight_if_empty(sv)

    return mapping


def prepare_pier_mapping(pier: dict) -> dict:
    p = pier or {}
    mapping = {}

    for k, v in p.items():
        if k == "uid":
            continue
        sv = "" if v is None else str(v)
        sv = _normalize_dash(sv)

        if k in FLOAT_1:
            sv = _fmt_float(sv, 1)
        elif k in FLOAT_2:
            sv = _fmt_float(sv, 2)

        mapping[f"{{{{pier.{k}}}}}"] = _keep_highlight_if_empty(sv)

    return mapping

# ==================================================
# DEFECTS TABLE (Форма 5)
# ==================================================

def fill_defects_table(doc: Document, defects: list):
    marker = "{{DEFECTS_TABLE}}"
    table = find_table_after_marker(doc, marker)

    if table is None:
        raise ValueError(
            "Не найдена таблица дефектов. "
            "Поставь маркер {{DEFECTS_TABLE}} перед нужной таблицей в шаблоне."
        )

    if len(table.columns) < 6:
        raise ValueError(
            f"Таблица дефектов найдена, но в ней {len(table.columns)} колонок. Нужно 6."
        )

    # очищаем строки после заголовка
    while len(table.rows) > 1:
        table._tbl.remove(table.rows[1]._tr)

    # группировка по разделам
    grouped = defaultdict(list)
    for rec in defects:
        grouped[rec.get("placement", "")].append(rec)

    placements_sorted = sorted(
        [p for p in grouped.keys() if p],
        key=lambda x: int(x.split(".", 1)[0]) if x.split(".", 1)[0].isdigit() else 999
    )

    counter = 1

    for placement in placements_sorted:
        entries = grouped[placement]

        # строка-заголовок раздела
        row = table.add_row()
        cell = row.cells[0]
        cell.merge(row.cells[-1])

        clean_name = placement.split(".", 1)[-1].strip()
        p = cell.paragraphs[0]
        p.text = ""
        run = p.add_run(clean_name)
        run.bold = True
        run.font.size = Pt(12)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        # строки дефектов
        for rec in entries:
            cells = table.add_row().cells
            if len(cells) < 6:
                raise ValueError(
                    "При добавлении строки таблицы получилось меньше 6 ячеек. "
                    "Проверь объединения ячеек в шаблоне."
                )

            cells[0].text = str(counter)
            cells[1].text = rec.get("location", "")
            cells[2].text = rec.get("name", "")
            cells[3].text = _format_units_for_docx(rec.get("option", ""))

            cats = []
            if rec.get("safety"):
                cats.append(f"Б{rec['safety']}")
            if rec.get("durability"):
                cats.append(f"Д{rec['durability']}")
            if rec.get("repairability"):
                cats.append(f"Р{rec['repairability']}")
            try:
                if rec.get("loadcap") and int(rec["loadcap"]) == 1:
                    cats.append("Г")
            except Exception:
                pass

            cells[4].text = ", ".join(cats)
            cells[5].text = rec.get("action", "")

            for c in cells:
                c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                for p in c.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            counter += 1

    remove_marker_everywhere(doc, marker)


# ==================================================
# EXPORT (BRIDGE PASSPORT)
# ==================================================

def export_to_docx(file_path: str, project: dict):
    """
    Экспорт отчёта в DOCX по шаблону.
    """
    if not isinstance(project, dict):
        raise TypeError("export_to_docx ожидает project=dict")

    shutil.copyfile(TEMPLATE_PATH, file_path)
    doc = Document(file_path)

    # ---------- Форма 1 (Основные сведения bridge.*) ----------
    bridge = project.get("bridge", {})
    mapping = prepare_bridge_mapping(bridge)
    replace_placeholders_everywhere(doc, mapping)

    # ---------- Форма 2 (Пролётные строения spans.*) ----------
    clone_block_between_markers(
        doc,
        start_marker="{{SPAN_FORM}}",
        end_marker="{{PIER_FORM}}",
        items=project.get("spans", []),
        prefix="span",
        mapping_builder=prepare_span_mapping
    )

# ---------- Форма 3 (Опоры piers.*) ----------
    clone_block_between_markers(
        doc,
        start_marker="{{PIER_FORM}}",
        end_marker="{{FORM4_START}}",
        items=project.get("piers", []),
        prefix="pier",
        mapping_builder=prepare_pier_mapping
    )
    # ---------- Форма 4  ----------
    replace_placeholders_everywhere(doc, {"{{FORM4_START}}": ""})
    
    # ---------- Форма 5 (дефекты defects.*) ----------
    fill_defects_table(doc, project.get("defects", []))


    # --- ФОТО ---
    photos = project.get("photos", {}) or {}
    folder = photos.get("folder", "") or ""

    cover = (photos.get("cover", {}) or {}).get("filename", "") or ""
    cover_path = os.path.join(folder, cover) if folder and cover else ""

    _insert_picture_at_marker(doc, "{{PHOTO_COVER}}", cover_path, w_cm=17.0, h_cm=12.0)


    gallery = photos.get("gallery", []) or []
    _fill_photos_gallery(doc, "{{PHOTOS_SECTION}}", gallery, folder, w_cm=16.0, h_cm=11.0)

    # --- SAVING ---
    doc.save(file_path)

# ==================================================
# EXPORT (TECHNICAL REPORT)
# ==================================================

def prepare_indexed_list_mapping(items: list, prefix: str) -> dict:
    """
    Делает mapping для плейсхолдеров вида:
      {{span0.key}}, {{span1.key}}, ...
      {{pier0.key}}, {{pier1.key}}, ...

    - items: список словарей (project["spans"] или project["piers"])
    - prefix: "span" или "pier"
    """
    mapping = {}

    if not items:
        return mapping

    for idx, item in enumerate(items):
        if not isinstance(item, dict):
            continue

        for k, v in item.items():
            if k == "uid":
                continue

            sv = "" if v is None else str(v)
            sv = _normalize_dash(sv)

            # форматирование чисел так же, как в паспорте
            if k in FLOAT_1:
                sv = _fmt_float(sv, 1)
            elif k in FLOAT_2:
                sv = _fmt_float(sv, 2)

            mapping[f"{{{{{prefix}{idx}.{k}}}}}"] = _keep_highlight_if_empty(sv)

    return mapping


def export_report_to_docx(file_path: str, project: dict):
    """
    Экспорт ТЕХНИЧЕСКОГО ОТЧЁТА в DOCX по шаблону inspection_report_template.docx.

    Использует те же данные:
    - bridge.* (Форма 1)
    - span0.*, span1.*, ... (по списку project["spans"])
    - pier0.*, pier1.*, ... (по списку project["piers"])
    - defects (таблица после {{DEFECTS_TABLE}})
    """
    if not isinstance(project, dict):
        raise TypeError("export_report_to_docx ожидает project=dict")

    # копируем шаблон отчёта
    shutil.copyfile(REPORT_TEMPLATE_PATH, file_path)
    doc = Document(file_path)

    # --- bridge.* ---
    bridge = project.get("bridge", {})
    replace_placeholders_everywhere(doc, prepare_bridge_mapping_report(bridge))

    # --- span0/span1/... ---
    spans = project.get("spans", [])
    replace_placeholders_everywhere(doc, prepare_indexed_list_mapping(spans, "span"))

    # --- pier0/pier1/... ---
    piers = project.get("piers", [])
    replace_placeholders_everywhere(doc, prepare_indexed_list_mapping(piers, "pier"))

    # --- ФОТО ---
    photos = project.get("photos", {}) or {}
    folder = photos.get("folder", "") or ""

    cover = (photos.get("cover", {}) or {}).get("filename", "") or ""
    cover_path = os.path.join(folder, cover) if folder and cover else ""

    _insert_picture_at_marker(doc, "{{PHOTO_COVER}}", cover_path, w_cm=17.0, h_cm=12.0)

    gallery = photos.get("gallery", []) or []
    _fill_photos_gallery(doc, "{{PHOTOS_SECTION}}", gallery, folder, w_cm=16.0, h_cm=11.0)


    # --- таблица дефектов ---
    fill_defects_table(doc, project.get("defects", []))

    doc.save(file_path)
